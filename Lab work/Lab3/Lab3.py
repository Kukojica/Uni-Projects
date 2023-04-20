import random
class BankAccount:
    def __init__(self, name, balance):
        self.name = name
        self.balance = balance
    def deposit(self, a):
        self.balance += a
        print(f"{a} было зачислено на аккаунт клиента {self.name}")
    def withdraw(self, a):
        if a > self.balance:
            print(f"Недостаточно средств на счете клиента {self.name}")
        else:
            self.balance -= a
            print(f"{a} было снято со счета клиента {self.name}")
    def check_balance(self):
        print(f"Текущий баланс средств на счете клиента {self.name} - {self.balance}")

acc_a = BankAccount('Anton', 5000)
print(acc_a.deposit(400))

class Car:
    def __init__(self, model):
        self.model = model
        self.speed = 0

    def start(self):
        self.speed += 10
        return (f'Текущая скорость - {self.speed} км/ч')

    def stop(self):
        self.speed = 0
        return (f'Текущая скорость - {self.speed} км/ч')
    def honk(self):
        return('Бип-бип')

car1 = Car('KIA')
print(car1.stop(), car1.honk())

class Person:
    def __init__(self, name, age, address):
        self.name = name
        self.age = age
        self.address = address

PersonA = Person('Anton', 22, 'Alexandr ave.')

print(PersonA.address)

class University:
    def __init__(self, name, address, students_n):
        self.name = name
        self. address = address
        self.students_n = students_n
YAGPU = University('ЯГПУ', 'Которосльная наб. 66', 200)
print(YAGPU.name)

class Rectangle:
    def __init__(self, width, height):
        self.width = width
        self.height = height
    def area(self):
        return self.width * self.height

    def p(self):
        return 2*(self.width + self.height)

r = Rectangle(20, 10)
print(r.area(), r.p())

class Product:
    def __init__(self, name:str,  price:float, brand='R'):
        self.name = name
        self.price = price
        self.brand = brand

class Cart:
    def __init__(self, owner:str):
        self.owner = owner
        self.l = list()

    def add(self, *args:Product):
        for arg in args:
            self.l.append(arg)

    def __str__(self):
        p_str=''
        for p in self.l:
            p_str += p.name + ', '
        return f'{self.owner}. Your products: {p_str}'

    def remove_one(self, p_name):
        for i in range(len(self.l)):
            if self.l[i].name == p_name:
                self.l.pop(i)
                break
    def counter(self):
        c=0
        for i in self.l:
            c+=1
        return f'Number of prods - {c}'
    def total(self):
        pr = 0
        for i in self.l:
            pr+= i.price
        return pr


mango = Product('Mango', 5.3)
milk = Product('Milk', 1.2)
apple = Product('Apple', 3.1)

p1 = Cart('Anton')


p1.add(mango, mango, milk, apple)
print(p1.remove_one('Mango'))
print(p1)

class Game:
    def __init__(self, p_name, score=0, hp=0, level=0):
        self.p_name = p_name
        self.score = score
        self.hp = hp
        self.level = level

    def start(self):
        self.hp += 3
        print(f'Welcome to the Game, {self.p_name}')

    def over(self):
        self.hp = 0
        self.level = 0
        self.score = 0

    def restart(self):
        if self.hp != 3:
            self.hp = 3
        if self.score != 0:
            self.score = 0
        if self.level != 0:
            self.level = 0
    def check_stats(self):
        if self.hp == 0:
            return 'Game Over'
        if self.score // 10:
            self.level += self.score // 10
        print(
            f'Имя игрока - {self.p_name}, кол-во жизней - {self.hp}, очки - {self.score}, уровень игрока - {self.level}')
    def get_score(self):
        self.score += random.randrange(0, 100, 10)
        print(f'Игрок {self.p_name} заработал очки - {self.score}')

p1 = Game('Kukoj')
p1.start()
p1.get_score()
p1.check_stats()

from docx import Document
from docx.shared import Cm
document = Document()
document.add_heading('Заголовок документа, созданного в Python', 0)
p = document.add_paragraph('Обычный абзац, в котором есть ')
p.add_run('полужирный шрифт').bold = True
p.add_run(' и немного ')
p.add_run('курсива.').italic = True
document.add_heading('Заголовок первого уровня', level=1)
document.add_paragraph('А прямо под ним подчеркнутый текст', style='Intense Quote')
document.add_heading('Куда я хочу поехать', level=2)
document.add_paragraph('Лондон ', style='List Bullet')
document.add_paragraph('Берлин ', style='List Bullet')
document.add_paragraph('Париж ', style='List Bullet')
document.add_picture('Л.jpg', width=Cm(6.5))
document.add_picture('Б.jpg', width=Cm(6.5))
document.add_picture('П.jpg', width=Cm(6.5))


document.save('demo.docx')

from openpyxl import Workbook
wb = Workbook()
ws = wb.active
ws.title = "Goods"
data = [['Название', 'Количество', 'Цена'], ['Карандаш', 5, 15], ['Ручка', 6, 29.9], ['Альбом', 2, 89], ['Тетрадь', 8, 16]]
for row in data:
    ws.append(row)
ws['D1'] = 'Общая стоимость'
ws['D2'] = '=B2:B5*C2:C5'
ws.formula_attributes['D2'] = {'t': 'array', 'ref': "D2:D5"}

wb.save('Wb.xlsx')